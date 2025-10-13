import os
import random
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from datetime import datetime

# ========== CONFIG ==========
BASE_DIR = r"D:\Assignments\Assignment2"
DOCX_DIR = os.path.join(BASE_DIR, "data", "word_docs")
PDF_DIR = os.path.join(BASE_DIR, "data", "word_pdfs")

TOTAL_DOCS = 5000
IMAGE_PERCENT = 0.30     # 30% documents will include an image
MULTIPAGE_FRACTION = 0.5 # 50% will be 2-3 pages (rest 1 page)

# Place your sample image here:
SAMPLE_IMAGE_PATH = os.path.join(BASE_DIR, "sample_image.jpg")

# For quick testing change TOTAL_DOCS = 20 etc.
# ============================

os.makedirs(DOCX_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

# Pre-select which docs are multi-page to ensure "half" exactly (or as close)
indices = list(range(1, TOTAL_DOCS + 1))
num_multi = int(TOTAL_DOCS * MULTIPAGE_FRACTION)
multi_indices = set(random.sample(indices, num_multi))

# Pre-select which docs will have images
num_with_images = int(TOTAL_DOCS * IMAGE_PERCENT)
image_indices = set(random.sample(indices, num_with_images))

def random_paragraph(min_words=70, max_words=120):
    words_pool = [
        "AI","data","forensics","pdf","binary","image","classifier","training","accuracy",
        "document","analysis","pattern","feature","histogram","byte","signature","model",
        "experiment","result","table","figure","method","report","section","algorithm",
        "random","distribution","scale","test","validation","metric","score","confusion"
    ]
    n = random.randint(min_words, max_words)
    return " ".join(random.choices(words_pool, k=n)) + "."

def build_doc(doc_id, pages, include_image=False):
    doc = Document()
    doc.add_heading(f"Document #{doc_id}", level=1)

    # Build content for each page
    for p in range(1, pages + 1):
        # Add a page-specific subheading
        doc.add_heading(f"Section {p}", level=2)
        # Add a few paragraphs per page
        for _ in range(random.randint(2, 4)):
            doc.add_paragraph(random_paragraph())

        # Optionally add an image on a random page (if include_image and this is chosen)
        if include_image and random.random() < 0.5:  # 50% chance to place image on this page
            if os.path.exists(SAMPLE_IMAGE_PATH):
                try:
                    doc.add_paragraph("Figure:")
                    doc.add_picture(SAMPLE_IMAGE_PATH, width=Inches(4.0))
                except Exception as e:
                    # If image insertion fails for any reason, keep going
                    print(f"⚠️ Warning: failed to insert image into doc {doc_id} page {p}: {e}")
            else:
                # image missing — this will be reported once for the missing file
                pass

        # Insert a page break unless it's the last page
        if p != pages:
            doc.add_page_break()

    return doc

# --- Generate DOCX files ---
start_time = datetime.now()
print(f"Start generating {TOTAL_DOCS} documents at {start_time.strftime('%Y-%m-%d %H:%M:%S')}")

for i in range(1, TOTAL_DOCS + 1):
    pages = 1
    if i in multi_indices:
        pages = random.choice([2, 3])  # 2 or 3 pages for selected multi-page docs

    include_image = (i in image_indices)

    doc = build_doc(i, pages, include_image=include_image)

    word_path = os.path.join(DOCX_DIR, f"doc_{i}.docx")
    try:
        doc.save(word_path)
    except Exception as e:
        print(f"⚠️ Error saving {word_path}: {e}")
    # progress log every 100 docs
    if i % 100 == 0:
        print(f"  Created {i}/{TOTAL_DOCS} .docx files...")

end_gen = datetime.now()
print(f"Completed DOCX generation at {end_gen.strftime('%Y-%m-%d %H:%M:%S')} (duration: {end_gen - start_time})")

# --- Convert to PDF using docx2pdf (requires MS Word on Windows) ---
print("Starting conversion of DOCX -> PDF using docx2pdf (this will use MS Word)...")
conv_start = datetime.now()
try:
    convert(DOCX_DIR, PDF_DIR)  # Converts entire directory
except Exception as e:
    print("❌ Conversion failed. Make sure Microsoft Word is installed and configured.")
    print("Exception:", e)
else:
    conv_end = datetime.now()
    print(f"Completed conversion at {conv_end.strftime('%Y-%m-%d %H:%M:%S')} (duration: {conv_end - conv_start})")
    print(f"PDFs saved in: {PDF_DIR}")

total_end = datetime.now()
print(f"All done. Total elapsed time: {total_end - start_time}")
